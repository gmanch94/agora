"""Generate Agora executive brief as a Word document (python-docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "artifacts/agora-executive-brief.docx"

# ---------------------------------------------------------------------------
# Palette  (RGB tuples → hex strings)
# ---------------------------------------------------------------------------
NAVY   = "1E3A5F"
STEEL  = "2D6A9F"
LIGHT  = "E8F0F8"
MUTED  = "64748B"
WHITE  = "FFFFFF"
BLACK  = "1E1E1E"
GREEN  = "166534"
GLIGHT = "DCFCE7"
AMBER  = "92400E"
ALIGHT = "FEF3C7"
GATE_BG = "F0F4F8"


def _rgb(hex6: str) -> RGBColor:
    r = int(hex6[0:2], 16)
    g = int(hex6[2:4], 16)
    b = int(hex6[4:6], 16)
    return RGBColor(r, g, b)


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex6: str):
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd if present
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    tcPr.append(shd)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set individual border sides on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for side, color in [("top", top), ("bottom", bottom),
                        ("left", left), ("right", right)]:
        if color is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
    tcPr.append(borders)


def _cell_para(cell) -> object:
    """Return the first paragraph in a cell (creates one if empty)."""
    if cell.paragraphs:
        return cell.paragraphs[0]
    return cell.add_paragraph()


def _no_table_borders(table):
    """Remove all borders from a table (used for callout box)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders_el = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        borders_el.append(el)
    tblPr.append(borders_el)


def _set_table_style(table):
    """Apply a minimal grid border style to a data table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Remove existing border spec
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders_el = OxmlElement("w:tblBorders")
    color = "B0BEC5"
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders_el.append(el)
    tblPr.append(borders_el)


def _para_space(para, before=0, after=0):
    """Set paragraph spacing in points (idempotent — removes any prior w:spacing)."""
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before * 20))  # twips
    spacing.set(qn("w:after"), str(after * 20))
    pPr.append(spacing)


# ---------------------------------------------------------------------------
# High-level helpers
# ---------------------------------------------------------------------------

def add_section_title(doc: Document, text: str):
    """Navy bold heading with steel underline border effect."""
    p = doc.add_paragraph()
    _para_space(p, before=14, after=2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = _rgb(NAVY)

    # Bottom border on the paragraph
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), STEEL)
    pb.append(bot)
    pPr.append(pb)


def add_body(doc: Document, text: str, indent_cm: float = 0.0):
    """Regular body paragraph."""
    p = doc.add_paragraph()
    _para_space(p, before=2, after=4)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = _rgb(BLACK)


def add_callout(doc: Document, text: str):
    """Single-cell table styled as a callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_table_borders(tbl)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, LIGHT)
    _set_cell_border(cell, top=STEEL, bottom=STEEL, left=STEEL, right=STEEL)

    para = _cell_para(cell)
    _para_space(para, before=6, after=6)
    para.paragraph_format.left_indent = Cm(0.4)
    para.paragraph_format.right_indent = Cm(0.4)
    run = para.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = _rgb(STEEL)

    # Spacer after
    spacer = doc.add_paragraph()
    _para_space(spacer, before=0, after=4)


def add_bullet(doc: Document, label: str, text: str):
    """Bold label + body text bullet."""
    p = doc.add_paragraph()
    _para_space(p, before=3, after=3)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)

    run_label = p.add_run(label + "  ")
    run_label.bold = True
    run_label.font.size = Pt(9.5)
    run_label.font.color.rgb = _rgb(NAVY)

    run_text = p.add_run(text)
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = _rgb(BLACK)


def add_table(doc: Document, rows: list, col_widths_in: list,
              status_col: int = -1, gate_col: int = -1, gate_val: str = ""):
    """
    n-column data table.
      rows[0]     = header labels
      col_widths_in = column widths in inches
      status_col  = column index whose text drives green/amber fill
      gate_col/gate_val = row flagged as a 'gate' step (italic steel style)
    """
    table = doc.add_table(rows=len(rows), cols=len(col_widths_in))
    table.autofit = False
    table.allow_autofit = False
    _set_table_style(table)

    # Set column widths (must come after autofit=False)
    for col_idx, width in enumerate(col_widths_in):
        for row in table.rows:
            row.cells[col_idx].width = Inches(width)

    for row_idx, row_data in enumerate(rows):
        tr = table.rows[row_idx]
        is_header = row_idx == 0
        is_gate = (not is_header and gate_col >= 0
                   and gate_val and gate_val in row_data[gate_col])
        is_stripe = (not is_header and not is_gate and row_idx % 2 == 1)

        for col_idx, text in enumerate(row_data):
            cell = tr.cells[col_idx]

            # Background fill
            if is_header:
                _set_cell_bg(cell, NAVY)
            elif is_gate:
                _set_cell_bg(cell, GATE_BG)
            elif is_stripe:
                _set_cell_bg(cell, LIGHT)
            else:
                _set_cell_bg(cell, WHITE)

            # Status column override
            status_bg = None
            status_text_color = None
            if not is_header and col_idx == status_col:
                val = text.lower()
                if val in ("implemented", "green", "written"):
                    status_bg = GLIGHT
                    status_text_color = GREEN
                elif val.strip():
                    status_bg = ALIGHT
                    status_text_color = AMBER
                if status_bg:
                    _set_cell_bg(cell, status_bg)

            para = _cell_para(cell)
            _para_space(para, before=3, after=3)
            para.paragraph_format.left_indent = Cm(0.15)
            para.paragraph_format.right_indent = Cm(0.1)

            run = para.add_run(text)
            run.font.size = Pt(9)

            if is_header:
                run.bold = True
                run.font.color.rgb = _rgb(WHITE)
            elif is_gate:
                run.italic = True
                run.font.color.rgb = _rgb(STEEL)
            elif status_text_color:
                run.bold = True
                run.font.color.rgb = _rgb(status_text_color)
            else:
                run.font.color.rgb = _rgb(BLACK)

    # Space after table
    spacer = doc.add_paragraph()
    _para_space(spacer, before=0, after=4)


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Cover block ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    _para_space(title_p, before=0, after=2)
    title_run = title_p.add_run("Agora")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = _rgb(NAVY)

    sub_p = doc.add_paragraph()
    _para_space(sub_p, before=0, after=2)
    sub_run = sub_p.add_run("Agentic Inter-Library Loan Orchestrator")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = _rgb(STEEL)

    meta_p = doc.add_paragraph()
    _para_space(meta_p, before=0, after=8)
    meta_run = meta_p.add_run("May 2026   |   Research Prototype   |   For internal review")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = _rgb(MUTED)

    add_callout(
        doc,
        "Agora automates decision-support for every ILL state transition "
        "while keeping a human staff click as the gate for every state change."
    )

    # ── Problem ──────────────────────────────────────────────────────────────
    add_section_title(doc, "Problem")
    add_body(doc,
        "Inter-Library Loan (ILL) is one of the most manual workflows in academic "
        "libraries. A single request typically demands 5-15 staff decisions spanning "
        "discovery, supplier selection, copyright clearance, shipment tracking, and "
        "reconciliation. Existing NCIP automation reduces steps by ~50% on the borrow "
        "side (NISO benchmarks) - but the remaining decisions are unstructured, untracked, "
        "and spread across disparate systems. Failures (item unavailable, supplier decline, "
        "lost in transit) are common, and legal compliance (CONTU, copyright) raises the "
        "bar for correctness."
    )

    # ── What Agora Is ────────────────────────────────────────────────────────
    add_section_title(doc, "What Agora Is")
    add_body(doc,
        "An agentic ILL orchestrator built on top of FOLIO/ReShare - the open-source "
        "library system stack already deployed across consortia. Agora adds a multi-agent "
        "advisory layer and a human-gated workflow engine over existing ISO 18626 "
        "infrastructure. It does not replace ReShare; it drives it."
    )

    # ── How It Works ─────────────────────────────────────────────────────────
    add_section_title(doc, "How It Works")
    add_table(
        doc,
        rows=[
            ("Stage", "Component", "What it does"),
            ("1 - Discovery",   "DiscoveryAgent",
             "Resolves citation / OpenURL to item holders via SRU catalog and CrossRef DOI lookup"),
            ("2 - Routing",     "RoutingAgent",
             "Ranks suppliers by SLA tier, reciprocity, lender load, proximity; "
             "optional Gemini LLM tie-breaker"),
            ("3 - Policy",      "PolicyAgent",
             "Checks CONTU copyright limits, patron eligibility, budget caps; flags hard violations"),
            ("→ Staff gate",    "Staff Console",
             "Staff reviews recommendation and rationale, then clicks Approve or Reject"),
            ("4 - Transaction", "TransactionAgent",
             "Drives ReShare over ISO 18626 (send request, ship, return) via async outbox worker"),
            ("5 - Tracking",    "TrackingAgent",
             "Monitors overdue loans, proposes recalls, flags unconfirmed receipts - three advisory tiers"),
            ("6 - Compensators","ReconciliationAgent",
             "Rolls back committed steps on failure: cancel, reroute, dispute resolution, override"),
        ],
        col_widths_in=[1.35, 1.55, 3.70],
        gate_col=1,
        gate_val="Staff",
    )
    add_body(doc,
        "Human-in-the-loop is non-negotiable. Every lifecycle state transition - "
        "Submit, Route, Approve (via an Approving hold), Ship, Receive, Return - requires "
        "an explicit staff approval. Agents advise; staff commits. This preserves legal "
        "and policy accountability."
    )

    # ── Key Design Decisions ─────────────────────────────────────────────────
    add_section_title(doc, "Key Design Decisions")
    add_table(
        doc,
        rows=[
            ("Decision", "Rationale"),
            ("Saga ledger (event-sourced)",
             "Every action is an immutable event. Full audit trail. "
             "Compensators undo exactly what was committed. Replay-safe by construction."),
            ("Outbox pattern (commit-then-enqueue)",
             "ISO 18626 messages to ReShare commit atomically with the ledger event - "
             "no silent drops, no double-delivery."),
            ("Advisory-only agents",
             "Legal and policy liability requires a human gate. Agents surface reasoning; "
             "staff own the decision. Enables FedRAMP alignment path."),
            ("Wraps ReShare, does not replace it",
             "ISO 18626 wire-level correctness stays in mod-rs. "
             "Agora drives it over REST. No XSD reimplementation."),
            ("LLM tie-breaking (optional)",
             "RoutingAgent uses deterministic rules by default. "
             "Opt-in Gemini Flash adapter resolves near-ties. "
             "95% top-1 accuracy on 20 labeled scenarios."),
        ],
        col_widths_in=[2.35, 4.25],
    )

    # ── Build Status ─────────────────────────────────────────────────────────
    add_section_title(doc, "Build Status")
    add_table(
        doc,
        rows=[
            ("Component", "Status"),
            ("7-state lifecycle (Submitted, Routed, Approving, Approved, Shipped, Received, Returned)",
             "Implemented"),
            ("Saga compensators for every forward step", "Implemented"),
            ("All 6 advisory agents (Discovery, Routing, Policy, Transaction, Tracking, Reconciliation)",
             "Implemented"),
            ("HTMX + Jinja2 staff console (inbox, detail, approve / reject / compensate / discover / override)",
             "Implemented"),
            ("Async outbox worker with Postgres multi-worker safety (SKIP LOCKED)",
             "Implemented"),
            ("NCIP fan-out (check-out on receive, check-in on return)", "Implemented"),
            ("Overdue scanner - 3 tiers: overdue / recall-proposed / unconfirmed-receipt",
             "Implemented"),
            ("ISO 18626 XSD validation harness", "Implemented"),
            ("LLM routing tie-breaker (Gemini 2.5 Flash via Vertex AI)", "Implemented"),
            ("Alembic migrations on real Postgres + CI gate", "Implemented"),
            ("247 automated tests (unit + property-based + end-to-end)", "Green"),
            ("15 Architecture Decision Records", "Written"),
        ],
        col_widths_in=[4.85, 1.75],
        status_col=1,
    )

    # ── Remaining Gaps ───────────────────────────────────────────────────────
    add_section_title(doc, "Remaining Gaps")
    add_table(
        doc,
        rows=[
            ("Item", "Blocker"),
            ("Real ReShare wire (live mod-rs tenant)",
             "Sandbox credentials needed"),
            ("Real NCIP HTTP/SOAP client",
             "Same blocker"),
            ("WorldCat holdings lookup",
             "OCLC sandbox key needed"),
            ("FedRAMP authorization",
             "Explicitly deferred - research prototype scope"),
            ("Patron-facing UI",
             "Out of scope for prototype"),
        ],
        col_widths_in=[3.35, 3.25],
        status_col=1,
    )

    # ── Prototype Validation ─────────────────────────────────────────────────
    add_section_title(doc, "Prototype Validation")
    add_bullet(doc, "End-to-end demo:",
        "make demo runs Submit to Return against an in-memory mock in under "
        "10 seconds, printing every ledger event.")
    add_bullet(doc, "Property-based tests:",
        "Arbitrary saga sequences including partial failures and replay - "
        "all compensators invert their forward step correctly (Hypothesis framework).")
    add_bullet(doc, "Routing eval harness:",
        "95% top-1 accuracy, 0.89 mean Spearman rank correlation against "
        "20 hand-labeled ILL scenarios using Gemini 2.5 Flash.")
    add_bullet(doc, "Postgres CI:",
        "Alembic round-trip (upgrade, downgrade, upgrade) and multi-worker outbox "
        "safety verified against postgres:15-alpine on every PR.")

    # ── Technology Stack ─────────────────────────────────────────────────────
    add_section_title(doc, "Technology Stack")
    add_table(
        doc,
        rows=[
            ("Layer", "Technology"),
            ("Language / Runtime",  "Python 3.11+  (built on 3.14.3)"),
            ("API Framework",       "FastAPI + pydantic v2"),
            ("Database",            "PostgreSQL 15  |  SQLAlchemy async  |  Alembic migrations"),
            ("UI",                  "HTMX 2.0.4 + Jinja2  (server-rendered, no build step)"),
            ("Agent LLM",           "Google Gemini 2.5 Flash via Vertex AI  (Google ADK)"),
            ("ILL Standards",       "ISO 18626  |  NCIP  |  SRU  |  OpenURL  (via FOLIO/ReShare mod-rs)"),
            ("Testing",             "pytest + pytest-asyncio + Hypothesis  (property-based)"),
            ("Quality Gates",       "ruff  |  mypy --strict  |  bandit  |  pip-audit  |  detect-secrets"),
        ],
        col_widths_in=[2.15, 4.45],
    )

    # ── Next Steps ───────────────────────────────────────────────────────────
    add_section_title(doc, "Next Steps to Production Path")
    for i, (label, text) in enumerate([
        ("Secure a ReShare sandbox tenant",
         "Unblocks the three remaining wired integrations: ReShare wire, "
         "NCIP client, WorldCat."),
        ("Pilot with one consortium member",
         "Validate routing quality against a real supplier population; "
         "collect approval latency data."),
        ("Auth layer",
         "HTTP Basic is in place; swap for institution SSO per ADR-0007 path."),
        ("FedRAMP readiness review",
         "Architecture is alignment-noted; controls not yet implemented."),
        ("Patron-facing UI",
         "Staff console is complete; patron submission form is out of scope "
         "for prototype but well-defined in PRD-05."),
    ], 1):
        add_bullet(doc, f"{i}.  {label}:", text)

    # ── Footer note ──────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    _para_space(footer_p, before=12, after=0)
    # Horizontal rule via paragraph bottom border
    pPr = footer_p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    top_el = OxmlElement("w:top")
    top_el.set(qn("w:val"), "single")
    top_el.set(qn("w:sz"), "4")
    top_el.set(qn("w:space"), "2")
    top_el.set(qn("w:color"), MUTED)
    pb.append(top_el)
    pPr.append(pb)

    footer_run = footer_p.add_run(
        "Full technical documentation: docs/ - PRDs (7), ADRs (15), architecture diagrams, "
        "runbook, solution design. Source code: src/agora/. "
        "247 automated tests. Repository: github.com/gmanch94/agora."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = _rgb(MUTED)

    doc.save(OUTPUT)
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    build()
